from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    document = Document()

    # Title
    title = document.add_heading('Chapter 5: System Development', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5.1
    document.add_heading('5.1 Navigation/Module Structure', level=2)
    document.add_paragraph('The proposed smart attendance system is divided into three main parts to make sure it covers everything needed. The first part is the Gate Unit, which is kept at the school entrance. The second part is the Bus Unit, which travels with the school bus. The third part is the Cloud Backend, which connects everything together using the internet.')

    document.add_heading('1. Gate Unit Module (Raspberry Pi)', level=3)
    document.add_paragraph('This module is placed at the main gate of the school. I decided to use a "Thick Client" architecture here because the gate unit needs to work even if the school internet is down. If the internet fails, we cannot stop students from entering, so the database is stored locally inside the Raspberry Pi.')
    document.add_paragraph('Input Sub-module: This part is responsible for reading the student ID cards. I used an RC522 RFID reader connected via the SPI interface to the Raspberry Pi.', style='List Bullet')
    document.add_paragraph('Logic Sub-module: This is the brain of the unit. It takes the ID from the scanner and checks it against a "students.json" file stored on the device. It also has a cool-down timer to stop students from accidentally scanning twice in one minute.', style='List Bullet')
    document.add_paragraph('Communication Sub-module: This part handles the SMS. It is connected to a generated GSM module that sends a text to the parent immediately. It also tries to sync with Google Sheets in the background when internet is available.', style='List Bullet')

    document.add_heading('2. Bus Unit Module (ESP32)', level=3)
    document.add_paragraph('For the bus, I needed something smaller and cheaper than a Raspberry Pi, so I used an ESP32. This is designed as a "Thin Client". This means it doesnt think that much on its own; it just asks the cloud if a card is valid or not. This makes the code simpler but it means it needs a good internet connection.')
    document.add_paragraph('Network Handler: This was the hardest part to setup. It uses a SIM800L module to access GPRS internet (using Dialog APN) because there is no WiFi on the road.', style='List Bullet')
    document.add_paragraph('Hardware Interface: This simply controls the RFID reader and beeps the buzzer. Green light means "Boarded" or "Left" successfully, and Red light means the card is unknown.', style='List Bullet')
    document.add_paragraph('Cloud Connector: This function sends a secure HTTPS request to my Google Script. It waits for the script to say "allowed" before letting the student pass.', style='List Bullet')

    document.add_heading('3. Cloud Module (Google Script & Sheets)', level=3)
    document.add_paragraph('This is where all the data is actually stored.')
    document.add_paragraph('Database: I used Google Sheets because it is free and easy to use. The school admin can just open the sheet and see who came to school today.', style='List Bullet')
    document.add_paragraph('API Layer: Since ESP32 cannot write to sheets directly easily, I wrote a small Google Apps Script. This script acts like a middleman. It takes the request, does the logic, and saves the data.', style='List Bullet')

    # 5.2
    document.add_heading('5.2 Development Environment', level=2)
    document.add_paragraph('To build this system, I had to use a mix of different hardware and software tools. Below is the list of what was used and why I chose them.')
    
    document.add_heading('Hardware Environment', level=3)
    document.add_paragraph('Raspberry Pi 3 Model B: I used this for the gate because it is powerful like a mini computer. It runs Linux, so it was easy to save the database files locally.', style='List Bullet')
    document.add_paragraph('ESP32 DevKit V1: This was selected for the bus unit. It is very low power and has built-in WiFi (though I mostly used GPRS). It is much cheaper than the Pi, which is good for making many units.', style='List Bullet')
    document.add_paragraph('RC522 RFID Reader: This is a standard affordable reader that works with 13.56 MHz cards. It connects easily to both Pi and ESP32 using SPI pins.', style='List Bullet')
    document.add_paragraph('SIM800L GSM Module: I used this module for mobile connectivity. It supports 2G/GPRS, which is enough for sending small text data and SMS messages.', style='List Bullet')

    document.add_heading('Software Environment', level=3)
    document.add_paragraph('Python 3.7: I used Python for the Raspberry Pi coding. It is very beginner-friendly and has great libraries like "gspread" for Google Sheets and "RPi.GPIO" for the hardware pins.', style='List Bullet')
    document.add_paragraph('Arduino IDE: This was used to write code for the ESP32. The language is C++, which is a bit harder than Python but gives more control over the hardware.', style='List Bullet')
    document.add_paragraph('Google Apps Script: This is based on JavaScript. I used it to connect the physical devices to the Google Sheet online. It runs on Google servers so I didnt need to buy a hosting plan.', style='List Bullet')

    # 5.3
    document.add_heading('5.3 Tools and Technologies', level=2)
    document.add_paragraph('I used several technologies to make the development process easier. Here is a summary of them:')
    
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Tool/Technology'
    hdr_cells[2].text = 'How it was used'
    
    data = [
        ('Language', 'Python', 'Used for the main logic on the Raspberry Pi Gate Unit.'),
        ('Language', 'C++ / Arduino', 'Used for programming the ESP32 microcontroller.'),
        ('Library', 'TinyGSM', 'A very useful library that handles the complex AT commands for the SIM800L module.'),
        ('Library', 'gspread', 'Python library used to read and write to Google Sheets from the Pi.'),
        ('Platform', 'Google Sheets', 'Functions as the main database for storing student details and attendance logs.')
    ]

    for cat, tool, usage in data:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = tool
        row_cells[2].text = usage

    # 5.4
    document.add_heading('5.4 Deployment Architecture', level=2)
    document.add_paragraph('The deployment follows a Hybrid Edge-Cloud Architecture.')
    
    # Insert Architecture Diagram
    try:
        document.add_picture('diagram_arch.png', width=Inches(6))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('Figure 5.4: Deployment Architecture', style='Caption')
    except Exception:
        document.add_paragraph('[Diagram: Architecture Diagram Missing]')

    document.add_paragraph('Edge Computing (Gate Unit): Critical processing happens locally on the Raspberry Pi. This ensures that students can still be scanned and parents notified via SMS even if the school internet goes down.', style='List Bullet')
    
    # Insert Gate Diagram
    try:
        document.add_picture('diagram_gate.png', width=Inches(6))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('Figure 5.1: Gate Unit Data Flow', style='Caption')
    except Exception:
        document.add_paragraph('[Diagram: Gate Unit Missing]')

    document.add_paragraph('Cloud Computing (Bus Unit): The Bus unit relies on the cloud for validation. It connects via the cellular network (Dialog GPRS) to the Google Script API.', style='List Bullet')

    # Insert Bus Diagram
    try:
        document.add_picture('diagram_bus.png', width=Inches(6))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('Figure 5.2: Bus Unit Data Flow', style='Caption')
    except Exception:
        document.add_paragraph('[Diagram: Bus Unit Missing]')

    # 5.5
    document.add_heading('5.5 Major Code Segments', level=2)
    
    document.add_heading('5.5.1 Cloud Synchronization (Raspberry Pi)', level=3)
    document.add_paragraph('This Python segment demonstrates the "Background Sync" feature. It ensures the local database is always up-to-date with the Google Sheet without blocking the main scanning loop.')
    document.add_paragraph("""def sync_students_from_sheet():
    global STUDENTS_CACHE
    try:
        # Auth with Google Sheets
        scope = ['https://spreadsheets.google.com/feeds', ...]
        creds = ServiceAccountCredentials.from_json_keyfile_name(CREDENTIALS_FILE, scope)
        client = gspread.authorize(creds)
        
        # Fetch all records
        sheet = client.open(SHEET_NAME).worksheet("Students")
        records = sheet.get_all_records()
        
        # Update in-memory cache and save to local JSON backup
        STUDENTS_CACHE = { str(r['UID']): r for r in records }
        with open(STUDENTS_FILE, 'w') as f:
            json.dump(STUDENTS_CACHE, f)
            
    except Exception as e:
        print(f"[CLOUD] Sync Failed: {e}")""", style='No Spacing')

    document.add_heading('5.5.2 GPRS Data Transmission (ESP32)', level=3)
    document.add_paragraph('This C++ segment shows how the ESP32 switches from a standard function call to a GPRS network request. It utilizes the TinyGSM library to establish a cellular data connection and query the API.')
    document.add_paragraph("""// From: SmartAttendanceESP32_GPRS.ino
void sendDataToScript(String uid, String type) {
  // Construct the secure API URL
  String path = SCRIPT_PATH + "?action=handleScan" + "&uid=" + uid + "&type=" + type;

  // Initiate Request via SIM800L Modem
  http.beginRequest();
  http.get(path);
  http.endRequest();

  // Parse Response
  String responseBody = http.responseBody();
  
  if (responseBody.indexOf("\\"allowed\\"") > 0) {
      triggerSuccess(); // Green LED
      // Extract parent phone and send SMS...
  } else {
      triggerFail(); // Red LED
  }
}""", style='No Spacing')

    document.add_heading('5.5.3 Cloud API Logic (Google Apps Script)', level=3)
    document.add_paragraph('This JavaScript function resides in the cloud. It acts as the central validator, receiving the UID from the bus, finding the corresponding student, and returning the parent\'s phone number.')
    document.add_paragraph("""// From: google_script_api.js
function handleRequest(e) {
    if (e.parameter.action == "handleScan") {
        var uid = e.parameter.uid;
        
        // Search for student in the "Students" sheet
        var student = findStudent(ss, uid);

        if (student) {
            // Log attendance to "Attendance" sheet
            ss.getSheetByName("Attendance").appendRow([new Date(), uid, student.name, ...]);
            
            // Return 'allowed' status and Phone Number to the device
            return response({
                status: "allowed",
                phone: student.phone
            });
        }
    }
}""", style='No Spacing')

    document.save('Chapter_5_System_Development_v3.docx')
    print("Document created successfully: Chapter_5_System_Development_v3.docx")

if __name__ == "__main__":
    create_doc()
